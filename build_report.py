"""
Generates Assignment3_Report.docx — Cloud Computing Assignment 3
Run: python3 build_report.py
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

BASE = os.path.dirname(os.path.abspath(__file__))
SS  = os.path.join(BASE, "Screenshots", "Assignment3")

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Helpers ───────────────────────────────────────────────────────────────────

DARK_BLUE  = RGBColor(0x0F, 0x34, 0x60)
MID_BLUE   = RGBColor(0x00, 0x78, 0xD4)
LIGHT_GREY = RGBColor(0xF4, 0xF6, 0xF9)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
TEXT_GREY  = RGBColor(0x55, 0x55, 0x55)


def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)


def para(text, size=11, bold=False, color=None, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.font.size  = Pt(size)
    run.font.bold  = bold
    if color:
        run.font.color.rgb = color
    return p


def heading(text, level=1):
    """Part heading with dark blue bar style."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.font.size  = Pt(16 if level == 1 else 13)
    run.font.bold  = True
    run.font.color.rgb = DARK_BLUE
    # bottom border
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    '6')
    bot.set(qn('w:space'), '4')
    bot.set(qn('w:color'), '0F3460')
    pBdr.append(bot)
    pPr.append(pBdr)
    return p


def sub_heading(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text.upper())
    run.font.size  = Pt(9)
    run.font.bold  = True
    run.font.color.rgb = TEXT_GREY
    return p


def add_screenshot(path, caption, width=6.0):
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(width))
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        para(f"[Screenshot not found: {os.path.basename(path)}]", color=RGBColor(0xAA, 0, 0))
    c = doc.add_paragraph(caption)
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c.paragraph_format.space_after = Pt(10)
    run = c.runs[0]
    run.font.size   = Pt(9)
    run.font.italic = True
    run.font.color.rgb = TEXT_GREY


def simple_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    hdr = t.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = WHITE
        cell.paragraphs[0].runs[0].font.size = Pt(10)
        set_cell_bg(cell, '0F3460')
    # Data rows
    for row_data in rows:
        row = t.add_row()
        for i, val in enumerate(row_data):
            cell = row.cells[i]
            cell.text = str(val)
            cell.paragraphs[0].runs[0].font.size = Pt(10)
    # Column widths
    if col_widths:
        for row in t.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph()  # spacing


# ═════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ═════════════════════════════════════════════════════════════════════════════

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(60)
p.paragraph_format.space_after  = Pt(4)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("CLOUD COMPUTING")
run.font.size = Pt(13)
run.font.bold = True
run.font.color.rgb = MID_BLUE

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(32)
run = p.add_run("Assignment 3 — Infrastructure Automation, Cost Analysis & Monitoring")
run.font.size = Pt(22)
run.font.bold = True
run.font.color.rgb = DARK_BLUE

para("ShopCloud — Azure E-Commerce Platform", size=14, align=WD_ALIGN_PARAGRAPH.CENTER, color=TEXT_GREY, space_after=48)

info = [
    ("Student",     "Pranav"),
    ("Submission",  "April 13, 2026"),
    ("Platform",    "Microsoft Azure"),
    ("IaC Tool",    "Terraform v1.5+"),
    ("App URL",     "https://shopcloud-pranav.azurewebsites.net"),
]
t = doc.add_table(rows=len(info), cols=2)
t.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (k, v) in enumerate(info):
    t.rows[i].cells[0].text = k
    t.rows[i].cells[1].text = v
    t.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
    t.rows[i].cells[0].paragraphs[0].runs[0].font.size = Pt(11)
    t.rows[i].cells[1].paragraphs[0].runs[0].font.size = Pt(11)
    t.rows[i].cells[0].width = Inches(1.8)
    t.rows[i].cells[1].width = Inches(4.2)

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
# PART 1 — TERRAFORM
# ═════════════════════════════════════════════════════════════════════════════

heading("Part 1 — Terraform Deployment  [4 marks]")

sub_heading("Overview")
para(
    "The ShopCloud infrastructure from Assignment 2 was fully codified using Terraform (HashiCorp "
    "Configuration Language). Since all Azure resources were originally provisioned through the Azure "
    "Portal, the terraform import workflow was used to bring all 14 existing resources under Terraform "
    "management without destroying or recreating any of them. Configuration files were then aligned with "
    "actual Azure resource settings to produce a clean plan with 0 resources to destroy."
)

sub_heading("Terraform File Structure")
simple_table(
    ["File", "Purpose"],
    [
        ("main.tf",          "Provider config — azurerm ~> 3.90, subscription ID"),
        ("variables.tf",     "All input variables with defaults and sensitive flags"),
        ("resource_group.tf","Azure Resource Group — rg-product-catalogue"),
        ("appservice.tf",    "App Service Plan (F1) + Linux Web App (Node 20 LTS)"),
        ("database.tf",      "PostgreSQL Flexible Server, database, 3 firewall rules"),
        ("storage.tf",       "Storage Account (Standard LRS) + Blob Container"),
        ("cognitive.tf",     "Computer Vision, Language AI, Translator (all F0 free)"),
        ("monitoring.tf",    "Application Insights (workspace-based)"),
        ("outputs.tf",       "App URL, DB FQDN, storage endpoint, AI keys (sensitive)"),
        ("import.sh",        "One-time script to import all 14 existing Azure resources"),
    ],
    col_widths=[1.8, 4.4]
)

sub_heading("Resources Managed by Terraform")
simple_table(
    ["Terraform Resource", "Azure Service", "SKU / Tier"],
    [
        ("azurerm_resource_group.main",                        "Resource Group",               "rg-product-catalogue"),
        ("azurerm_service_plan.main",                          "App Service Plan",             "F1 (Free, Linux)"),
        ("azurerm_linux_web_app.main",                         "Web App",                      "Node 20 LTS"),
        ("azurerm_postgresql_flexible_server.main",            "PostgreSQL Flexible Server",   "B_Standard_B1ms, 32 GB"),
        ("azurerm_postgresql_flexible_server_database.main",   "PostgreSQL Database",          "productcatalogue"),
        ("azurerm_postgresql_flexible_server_firewall_rule ×3","Firewall Rules",               "Azure services + dev IPs"),
        ("azurerm_storage_account.main",                       "Blob Storage Account",         "Standard LRS"),
        ("azurerm_storage_container.products",                 "Blob Container",               "product-images"),
        ("azurerm_cognitive_account.vision",                   "Computer Vision",              "F0 (Free)"),
        ("azurerm_cognitive_account.language",                 "Language AI",                  "F0 (Free)"),
        ("azurerm_cognitive_account.translator",               "Translator",                   "F0 (Free)"),
        ("azurerm_application_insights.main",                  "Application Insights",         "Workspace-based"),
    ],
    col_widths=[2.6, 2.2, 1.4]
)

# terraform validate
sub_heading("terraform validate")
para(
    "All 9 .tf configuration files were validated for syntax and internal consistency. "
    "Terraform reported: Success! The configuration is valid."
)
validate_path = os.path.join(SS, "terraform_validate.png")
add_screenshot(validate_path, "terraform validate — Success! The configuration is valid.")

# terraform plan
sub_heading("terraform plan")
para(
    "The plan confirmed 0 resources to add, 2 to change, 0 to destroy. The 2 in-place "
    "updates were: (1) syncing app settings sensitive values on the Web App, and (2) re-applying "
    "the PostgreSQL administrator password — unavoidable after import because Azure never exposes "
    "passwords through its API. No resources were recreated or destroyed."
)
add_screenshot(os.path.join(SS, "terraform_plan1.png"),
               "terraform plan (1/2) — State refresh confirming all 14 resources in Azure")
add_screenshot(os.path.join(SS, "terraform_plan2.png"),
               "terraform plan (2/2) — Plan: 0 to add · 2 to change · 0 to destroy")

# terraform apply
sub_heading("terraform apply")
para(
    "After reviewing and confirming the plan, terraform apply was executed. Both in-place updates "
    "completed successfully in under 60 seconds. "
    "Final result: Apply complete! Resources: 0 added, 2 changed, 0 destroyed."
)
add_screenshot(os.path.join(SS, "terraform_apply.png"),
               "terraform apply — Apply complete, all outputs confirmed (app URL, DB FQDN, storage endpoint, AI endpoints)")

# terraform show
sub_heading("terraform show")
para(
    "terraform show displays the complete current state of all 14 managed resources, confirming "
    "Terraform has full visibility of the deployed infrastructure. Selected pages shown below."
)

show_files = [
    ("terraform_show1.png", "Application Insights + Cognitive Services (Language, Translator)"),
    ("terraform_show2.png", "Cognitive Services (Vision) + Storage Account"),
    ("terraform_show3.png", "PostgreSQL Flexible Server configuration"),
    ("terraform_show4.png", "PostgreSQL Firewall Rules"),
    ("terraform_show5.png", "App Service Plan + Web App site config"),
    ("terraform_show6.png", "Web App app_settings (sensitive values)"),
    ("terraform_show7.png", "Web App outputs — app URL and DB FQDN"),
    ("terraform_show8.png", "Storage outputs — blob endpoint"),
    ("terraform_show9.png", "AI endpoints and App Insights instrumentation key"),
]
for fname, caption in show_files:
    add_screenshot(os.path.join(SS, fname), f"terraform show — {caption}", width=5.8)

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
# PART 2 — COST ESTIMATION
# ═════════════════════════════════════════════════════════════════════════════

heading("Part 2 — Cost Estimation  [1 mark]")

sub_heading("Assumptions")
simple_table(
    ["Parameter", "Value", "Reasoning"],
    [
        ("Number of users",    "100 / month",       "Small demo deployment — instructors, TAs, and testers"),
        ("Requests per day",   "500 / day",         "~5 requests per active user per day"),
        ("Compute",            "730 hours / month", "Single App Service instance running continuously"),
        ("Blob storage",       "1 GB",              "Product image catalogue on Azure Blob Storage"),
        ("Database queries",   "~15,000 / month",   "500 req/day × 30 days, mix of reads and writes"),
        ("AI API calls",       "< 5,000 / month",   "Within free tier limits for Vision, Language, Translator"),
        ("Log ingestion",      "< 5 GB / month",    "Within Azure Monitor / App Insights free tier"),
    ],
    col_widths=[1.8, 1.6, 2.8]
)

sub_heading("Monthly Cost Breakdown")
simple_table(
    ["Service", "Tier", "Monthly Cost (INR)"],
    [
        ("App Service (Web App)",          "F1 Free",         "₹0.00"),
        ("Azure Database for PostgreSQL",  "B1ms Burstable",  "₹381.10"),
        ("Storage Account (Blob LRS)",     "Standard LRS",    "₹24.08"),
        ("Computer Vision",                "F0 Free",         "₹0.00"),
        ("Language AI",                    "F0 Free",         "₹0.00"),
        ("Translator",                     "F0 Free",         "₹0.00"),
        ("Azure Monitor / App Insights",   "5 GB Free",       "₹0.00"),
        ("TOTAL",                          "—",               "₹405.18"),
    ],
    col_widths=[2.6, 2.0, 1.6]
)

para(
    "The dominant cost is the PostgreSQL Flexible Server (B1ms) at ₹381.10/month (~94% of total). "
    "All compute, AI, and monitoring services fall within Azure's free tiers at this usage scale. "
    "At production scale (10,000+ users), the App Service would need upgrading to at least B2 "
    "(~₹1,200/month) and PostgreSQL to GP_Standard_D2s_v3 (~₹3,500/month)."
)

sub_heading("Azure Pricing Calculator Screenshot")
add_screenshot(os.path.join(SS, "Estimate.png"),
               "Azure Pricing Calculator — Estimated monthly cost: ₹405.18  |  Upfront: ₹0.00")

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
# PART 3 — MONITORING
# ═════════════════════════════════════════════════════════════════════════════

heading("Part 3 — Monitoring  [3 marks]")

sub_heading("Traffic Simulation — Locust Setup")
para(
    "Locust 2.43.4 was used to simulate realistic e-commerce user behaviour against the live Azure "
    "deployment. The test spawned 20 concurrent users at a ramp-up rate of 2 users/second and ran "
    "for approximately 10 minutes."
)
simple_table(
    ["Parameter", "Value"],
    [
        ("Tool",             "Locust 2.43.4 (Python)"),
        ("Target",           "https://shopcloud-pranav.azurewebsites.net"),
        ("Concurrent users", "20"),
        ("Spawn rate",       "2 users/second"),
        ("Total requests",   "3,999"),
        ("Failure rate",     "0%"),
        ("Endpoints tested", "11 (browse, search, filter, detail, reviews, chat, health)"),
    ],
    col_widths=[2.2, 4.0]
)

sub_heading("Task Weighting")
para(
    "Tasks were weighted to reflect realistic e-commerce usage — reads dominate over writes:"
)
simple_table(
    ["Endpoint", "Weight", "Reason"],
    [
        ("GET /api/products (browse)",     "10×", "Most common action — catalogue browsing"),
        ("GET /api/products (category)",   "6×",  "Filter by category — common navigation"),
        ("GET /api/products (search)",     "5×",  "Keyword search"),
        ("GET /api/products/:id",          "5×",  "Product detail page"),
        ("GET /api/products/categories",   "4×",  "Navbar category list"),
        ("GET /api/products (featured)",   "3×",  "Homepage featured products"),
        ("GET /api/products/stats",        "3×",  "Homepage stats widget"),
        ("GET /api/products (sorted)",     "2×",  "Sort by price / rating"),
        ("POST /api/products/:id/reviews", "2×",  "Submit review — triggers Language AI"),
        ("POST /api/products/:id/chat",    "2×",  "AI chat — triggers Language AI"),
        ("GET /api/health",                "2×",  "Health check / monitoring ping"),
    ],
    col_widths=[2.8, 0.8, 2.6]
)

sub_heading("i) Performance Metrics Screenshots")

add_screenshot(os.path.join(SS, "locust_stats_table.png"),
               "Locust Statistics Table — 3,999 requests · 0 failures · median 152ms · max 2,189ms")
add_screenshot(os.path.join(SS, "locust_stats_chart.png"),
               "Locust Charts — ~6–8 RPS sustained · 95th percentile stable ~800ms · brief spike at ramp-up")
add_screenshot(os.path.join(SS, "insight_fav.png"),
               "Azure Application Insights Overview — 4,276 requests · avg server response 3.42ms · 0 failed requests · 100% availability")
add_screenshot(os.path.join(SS, "insight_server.png"),
               "Azure Application Insights Performance — POST /reviews (184ms) and POST /chat (124ms) are the slowest operations")

sub_heading("ii) Insights & Analysis")

insights = [
    ("Zero failures under load",
     "All 3,999 requests completed successfully with a 0% error rate across all 11 endpoints, "
     "demonstrating the stability of the Azure App Service and PostgreSQL backend under 20 concurrent users."),

    ("AI endpoints are the performance bottleneck",
     "POST /reviews (184ms avg) and POST /chat (124ms avg) were the two slowest operations because they "
     "make synchronous calls to Azure Language AI for sentiment analysis and key phrase extraction. "
     "All read-only endpoints responded in under 100ms on average."),

    ("App Service F1 cold-start effect",
     "The F1 free tier does not support always-on. The brief response time spike visible in the Locust "
     "chart at approximately 4:28 PM corresponds to a cold-start event where the app woke from idle. "
     "After warm-up (~30 seconds), response times stabilised around 120–160ms."),

    ("Azure-side vs. client-side latency gap",
     "Azure Application Insights recorded an average server response of just 3.42ms, while Locust "
     "showed a ~152ms median client-side. The ~150ms difference represents network transit time between "
     "the local test machine in India and the East Asia Azure region — not server processing overhead."),

    ("PostgreSQL handles concurrent load well",
     "The B_Standard_B1ms database tier handled all concurrent read/write traffic without any "
     "connection timeouts or query errors, confirming it is appropriately sized for this workload."),

    ("Scalability recommendation",
     "For 500+ concurrent users, the App Service should be upgraded to at least B2 (2 vCores) with "
     "always-on enabled, and PostgreSQL to GP_Standard_D2s_v3 to handle the increased AI endpoint "
     "concurrency without response time degradation. Caching frequent catalogue queries (e.g. via "
     "Azure Cache for Redis) would also significantly reduce database load."),
]

for title, body in insights:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(8)
    run_title = p.add_run(title + ": ")
    run_title.font.bold = True
    run_title.font.size = Pt(11)
    run_title.font.color.rgb = DARK_BLUE
    run_body = p.add_run(body)
    run_body.font.size = Pt(11)

# ── Save ──────────────────────────────────────────────────────────────────────
out = os.path.join(BASE, "Assignment3_Report.docx")
doc.save(out)
print(f"✅  Saved: {out}")
